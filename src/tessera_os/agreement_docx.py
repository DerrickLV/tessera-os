"""Render an assembled agreement or a structure recommendation as branded Word.

Entirely in Python -- no Node, no external service, no network.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .clauses import AssembledDraft

NAVY = "0B1F3A"
GOLD = "B8963E"
GOLD_DEEP = "96792F"
INK = "1B2430"
MUTED = "5B6472"
LINE = "E2E5EA"

_BOLD = re.compile(r"\*\*(.+?)\*\*")
# Bold first, then single-asterisk italics. Matched together so ``**x** *y*``
# does not leave stray asterisks on the page -- which is what happened to the
# "Confirm:" notes in the first structuring memo.
_EMPHASIS = re.compile(r"\*\*(?P<bold>.+?)\*\*|\*(?P<italic>[^*\n]+?)\*")


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _runs(text: str) -> list[dict[str, str | bool]]:
    """Split a Markdown line into bold, italic, and plain runs."""
    runs: list[dict[str, str | bool]] = []
    cursor = 0
    for match in _EMPHASIS.finditer(text):
        if match.start() > cursor:
            runs.append({"text": text[cursor:match.start()], "bold": False})
        if match.group("bold") is not None:
            runs.append({"text": match.group("bold"), "bold": True})
        else:
            runs.append({"text": match.group("italic"), "bold": False, "italic": True})
        cursor = match.end()
    if cursor < len(text):
        runs.append({"text": text[cursor:], "bold": False})
    return runs or [{"text": "", "bold": False}]


def _plain(text: str) -> str:
    """Strip emphasis markers. A heading is styled by its level, not by asterisks."""
    return _EMPHASIS.sub(lambda m: m.group("bold") or m.group("italic") or "", text)


def _blocks(markdown: str) -> list[dict]:
    """Flatten assembled Markdown into typed blocks for Word rendering."""
    blocks: list[dict] = []
    table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table
        if table:
            blocks.append({"kind": "table", "rows": table})
            table = []

    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(set(cell) <= {"-", ":", " "} for cell in cells):
                continue
            table.append(cells)
            continue
        flush_table()
        if not line:
            continue
        if line.startswith("# "):
            blocks.append({"kind": "title", "text": _plain(line[2:])})
        elif line.startswith("### "):
            blocks.append({"kind": "h3", "text": _plain(line[4:])})
        elif line.startswith("## "):
            blocks.append({"kind": "h2", "text": _plain(line[3:])})
        elif line == "---":
            blocks.append({"kind": "rule"})
        elif line.startswith("- "):
            blocks.append({"kind": "bullet", "runs": _runs(line[2:])})
        elif line.startswith("> "):
            blocks.append({"kind": "note", "text": _BOLD.sub(r"\1", line[2:])})
        else:
            blocks.append({"kind": "body", "runs": _runs(line)})
    flush_table()
    return blocks


def _set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def _set_paragraph_border(paragraph, *, side: str, color: str, size: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "4")
    border.set(qn("w:color"), color)
    borders.append(border)


def _add_runs(paragraph, runs: list[dict[str, str | bool]], *,
              size: float = 10, color: str = INK) -> None:
    for item in runs:
        run = paragraph.add_run(str(item["text"]))
        run.bold = bool(item["bold"])
        run.italic = bool(item.get("italic"))
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.color.rgb = _rgb(color)


def _add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = _rgb(MUTED)


def _configure_document(document: Document, *, banner: str, footer_label: str) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.font.color.rgb = _rgb(INK)

    header = section.header.paragraphs[0]
    brand = header.add_run("TESSERA GROUP")
    brand.bold = True
    brand.font.name = "Georgia"
    brand.font.size = Pt(8)
    brand.font.color.rgb = _rgb(NAVY)
    warning = header.add_run(f"   |   {banner}")
    warning.italic = True
    warning.font.name = "Calibri"
    warning.font.size = Pt(8)
    warning.font.color.rgb = _rgb(MUTED)
    _set_paragraph_border(header, side="bottom", color=GOLD, size="6")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = footer.add_run(f"{footer_label}   |   Page ")
    prefix.font.name = "Calibri"
    prefix.font.size = Pt(8)
    prefix.font.color.rgb = _rgb(MUTED)
    _add_page_field(footer, "PAGE")
    suffix = footer.add_run(" of ")
    suffix.font.name = "Calibri"
    suffix.font.size = Pt(8)
    suffix.font.color.rgb = _rgb(MUTED)
    _add_page_field(footer, "NUMPAGES")


def _render_blocks(document: Document, markdown: str) -> None:
    for block in _blocks(markdown):
        kind = block["kind"]
        if kind == "title":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(3)
            run = paragraph.add_run(block["text"])
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(17)
            run.font.color.rgb = _rgb(NAVY)
        elif kind == "h2":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(15)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run(block["text"])
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(12)
            run.font.color.rgb = _rgb(NAVY)
            _set_paragraph_border(paragraph, side="bottom", color=GOLD, size="4")
        elif kind == "h3":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(11)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(block["text"])
            run.bold = True
            run.font.name = "Georgia"
            run.font.size = Pt(10.5)
            run.font.color.rgb = _rgb(GOLD_DEEP)
        elif kind == "rule":
            paragraph = document.add_paragraph()
            _set_paragraph_border(paragraph, side="bottom", color=GOLD, size="6")
        elif kind == "bullet":
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            _add_runs(paragraph, block["runs"])
        elif kind == "note":
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.space_after = Pt(7)
            run = paragraph.add_run(block["text"])
            run.italic = True
            run.font.name = "Calibri"
            run.font.size = Pt(9.5)
            run.font.color.rgb = _rgb(GOLD_DEEP)
            _set_paragraph_border(paragraph, side="left", color=GOLD, size="12")
        elif kind == "table":
            rows = block["rows"]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            table.autofit = False
            # Equal columns put a one-character "#" beside a sentence and read as
            # a mistake. Weight each column by the longest cell in it, floored so
            # a narrow column stays legible and capped so one does not run away.
            longest = [max(len(row[col]) for row in rows) for col in range(len(rows[0]))]
            weights = [min(max(value, 6), 60) for value in longest]
            usable = Inches(7.0)
            for col, weight in enumerate(weights):
                width = int(usable * weight / sum(weights))
                table.columns[col].width = width
                for row in table.rows:
                    row.cells[col].width = width
            for row_index, row in enumerate(rows):
                for column_index, value in enumerate(row):
                    cell = table.cell(row_index, column_index)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    _add_runs(paragraph, [{"text": _BOLD.sub(r"\1", value),
                                           "bold": row_index == 0}],
                              size=9.5, color="FFFFFF" if row_index == 0 else INK)
                    if row_index == 0:
                        _set_cell_shading(cell, NAVY)
            document.add_paragraph().paragraph_format.space_after = Pt(4)
        else:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(7)
            paragraph.paragraph_format.line_spacing = 1.15
            _add_runs(paragraph, block["runs"])


class DocxRenderError(RuntimeError):
    """Raised when the Word renderer fails."""


def render_agreement_docx(draft: AssembledDraft, output: Path | str, *,
                          markdown: str | None = None) -> Path:
    """Write an assembled agreement to ``output`` as a branded DOCX file."""
    label = draft.profile.agreement_type.replace("_", " ").title()
    return _render(
        markdown if markdown is not None else draft.to_markdown(), output,
        banner="Draft — for counsel review",
        footer_label=f"{label} — {draft.profile.opportunity}",
        title=draft.profile.opportunity,
        subject="Draft agreement for qualified counsel review")


def render_structure_docx(recommendation, output: Path | str) -> Path:
    """Write a structure recommendation to ``output`` as a branded DOCX file.

    The memo is what the founder reads and what counsel works from, so it gets
    the same paper as the agreement it precedes -- and a banner that says what
    it is, because a structuring memo mistaken for an opinion is worse than one
    nobody reads.
    """
    venture = recommendation.profile.venture
    return _render(
        recommendation.to_markdown(), output,
        banner="Structural advice — not legal or tax advice",
        footer_label=f"Structure Recommendation — {venture}",
        title=venture,
        subject="Entity structure recommendation for qualified counsel review")


def _render(markdown: str, output: Path | str, *, banner: str, footer_label: str,
            title: str, subject: str) -> Path:
    target = Path(output)
    try:
        target.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        _configure_document(document, banner=banner, footer_label=footer_label)
        _render_blocks(document, markdown)
        document.core_properties.title = title
        document.core_properties.subject = subject
        document.save(target)
    except Exception as exc:
        raise DocxRenderError(f"Word rendering failed: {exc}") from exc
    if not target.is_file():
        raise DocxRenderError("Word rendering produced no file")
    return target
