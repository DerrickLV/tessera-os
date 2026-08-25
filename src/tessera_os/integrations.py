"""Read-only Microsoft Graph and SharePoint integration boundary."""

import json
import logging
import os
import time
from collections.abc import Callable, Iterator
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any
from urllib.error import HTTPError
from urllib.parse import quote, urlencode
from urllib.request import Request, urlopen

from docx import Document as DocxDocument

from .schemas import SourceDocument, UserContext

logger = logging.getLogger(__name__)


class IntegrationError(RuntimeError):
    pass


class GraphThrottleError(IntegrationError):
    def __init__(self, retry_after: float) -> None:
        super().__init__("Microsoft Graph throttled the read request")
        self.retry_after = retry_after


class SharePointPathNotFoundError(IntegrationError):
    """Raised when a configured ``root_path`` does not resolve to a folder.

    Carries the path so the caller can name it in the error surfaced to a
    partner, rather than the reader silently returning an empty list -- which
    is indistinguishable from "the folder is empty" and is exactly the defect
    this phase exists to fix.
    """

    def __init__(self, path: str) -> None:
        super().__init__(f"SharePoint path not found: {path!r}")
        self.path = path


@dataclass(frozen=True)
class GraphPage:
    values: list[dict[str, Any]]
    next_link: str | None


# Folder name (first path segment below ``root_path``) -> Lifecycle. Any other
# folder, or a file directly under root_path, is "source" -- evidence is the
# safe default; it grants no standing. See docs/BUILD_BRIEF_PHASE_2_LIBRARY_READING.md D2.
_LIFECYCLE_FOLDERS: dict[str, str] = {
    "Approved": "approved",
    "Drafts": "draft",
    "Source": "source",
}

# Narrow on purpose (D3): anything else lists with empty content and an honest
# ``content_available: False`` marker rather than guessing at a parser.
_SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md"}

_MAX_FETCH_BYTES = 10 * 1024 * 1024

_MAX_WALK_DEPTH = 5


def _extension(name: str) -> str:
    return os.path.splitext(name)[1].lower()


def _extract_text(extension: str, raw: bytes) -> str:
    if extension == ".docx":
        document = DocxDocument(BytesIO(raw))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    return raw.decode("utf-8")


def _lifecycle_for(folder_path: str) -> str:
    if not folder_path:
        return "source"
    first_segment = folder_path.split("/", 1)[0]
    return _LIFECYCLE_FOLDERS.get(first_segment, "source")


class MicrosoftGraphReader:
    """Delegated-token Graph client exposing GET requests only."""

    base_url = "https://graph.microsoft.com/v1.0"

    def __init__(self, token_provider: Callable[[], str], *,
                 transport: Callable[[str, dict[str, str]], dict[str, Any]] | None = None,
                 content_transport: Callable[[str, dict[str, str]], bytes] | None = None,
                 max_retries: int = 2,
                 sleeper: Callable[[float], None] = time.sleep) -> None:
        self._token_provider = token_provider
        self._transport = transport or self._get_json
        self._content_transport = content_transport or self._get_bytes
        self._max_retries = max_retries
        self._sleeper = sleeper

    @staticmethod
    def _get_json(url: str, headers: dict[str, str]) -> dict[str, Any]:
        request = Request(url, headers=headers, method="GET")
        try:
            with urlopen(request, timeout=30) as response:
                return json.loads(response.read())
        except HTTPError as exc:
            if exc.code == 429:
                retry_after = exc.headers.get("Retry-After", "1")
                try:
                    delay = max(0.0, float(retry_after))
                except ValueError:
                    delay = 1.0
                raise GraphThrottleError(delay) from exc
            raise IntegrationError("Microsoft Graph read failed") from exc
        except Exception as exc:
            raise IntegrationError("Microsoft Graph read failed") from exc

    @staticmethod
    def _get_bytes(url: str, headers: dict[str, str]) -> bytes:
        request = Request(url, headers=headers, method="GET")
        try:
            with urlopen(request, timeout=30) as response:
                return response.read()
        except HTTPError as exc:
            raise IntegrationError("Microsoft Graph content read failed") from exc
        except Exception as exc:
            raise IntegrationError("Microsoft Graph content read failed") from exc

    def _pages(self, path: str, params: dict[str, str] | None = None) -> Iterator[GraphPage]:
        url = path if path.startswith(self.base_url) else f"{self.base_url}/{path.lstrip('/')}"
        if params:
            url = f"{url}?{urlencode(params)}"
        headers = {"Authorization": f"Bearer {self._token_provider()}", "Accept": "application/json"}
        while url:
            if not url.startswith(f"{self.base_url}/"):
                raise IntegrationError("Graph pagination escaped the approved origin")
            for attempt in range(self._max_retries + 1):
                try:
                    payload = self._transport(url, headers)
                    break
                except GraphThrottleError as exc:
                    if attempt >= self._max_retries:
                        raise
                    self._sleeper(exc.retry_after)
            yield GraphPage(payload.get("value", []), payload.get("@odata.nextLink"))
            url = payload.get("@odata.nextLink")

    def _fetch_bytes(self, path: str) -> bytes:
        url = f"{self.base_url}/{path.lstrip('/')}"
        headers = {"Authorization": f"Bearer {self._token_provider()}",
                   "Accept": "application/octet-stream"}
        return self._content_transport(url, headers)

    def _resolve_root_path(self, *, site_id: str, drive_id: str, root_path: str) -> str:
        """Resolve a project's configured folder path to a Graph item reference.

        D1: the folder path *is* the project scope, and a misconfigured path
        must raise rather than silently return an empty listing -- the failure
        this whole phase exists to fix looked exactly like "the library is
        empty" from every layer above the reader.
        """
        encoded = quote(root_path, safe="/")
        url = f"{self.base_url}/sites/{site_id}/drives/{drive_id}/root:/{encoded}:"
        headers = {"Authorization": f"Bearer {self._token_provider()}", "Accept": "application/json"}
        try:
            payload = self._transport(url, headers)
        except IntegrationError as exc:
            raise SharePointPathNotFoundError(root_path) from exc
        item_id = payload.get("id")
        if not item_id or "folder" not in payload:
            raise SharePointPathNotFoundError(root_path)
        return f"items/{item_id}"

    def _walk(self, *, site_id: str, drive_id: str, item_path: str, folder_path: str,
              depth: int) -> Iterator[tuple[dict[str, Any], str]]:
        """Depth-bounded recursive descent under ``item_path``.

        ``depth`` is the depth of the items this call's ``children`` request
        returns (1 == directly under root_path). A folder is only descended
        into while doing so keeps its children at depth <= 5; a folder beyond
        that is logged and skipped, never a file -- a file at depth 5 is still
        returned.
        """
        path = f"sites/{site_id}/drives/{drive_id}/{item_path}/children"
        params = {"$select": "id,name,webUrl,lastModifiedDateTime,file,folder,size"}
        for page in self._pages(path, params):
            for item in page.values:
                name = item["name"]
                if "folder" in item:
                    child_folder_path = f"{folder_path}/{name}" if folder_path else name
                    if depth >= _MAX_WALK_DEPTH:
                        logger.warning(
                            "SharePoint walk exceeded max depth %d; skipping %s",
                            _MAX_WALK_DEPTH, child_folder_path)
                        continue
                    yield from self._walk(
                        site_id=site_id, drive_id=drive_id, item_path=f"items/{item['id']}",
                        folder_path=child_folder_path, depth=depth + 1)
                elif "file" in item:
                    yield item, folder_path

    def _build_document(self, *, item: dict[str, Any], site_id: str, drive_id: str,
                        context: UserContext, project_id: str, folder_path: str) -> SourceDocument:
        name = item["name"]
        extension = _extension(name)
        metadata: dict[str, Any] = {
            "site_id": site_id, "drive_id": drive_id,
            "lifecycle": _lifecycle_for(folder_path), "folder_path": folder_path,
        }
        content = ""
        size = item.get("size", 0)
        if extension not in _SUPPORTED_EXTENSIONS:
            metadata["content_available"] = False
        elif size > _MAX_FETCH_BYTES:
            metadata["content_available"] = False
            metadata["skipped_reason"] = "size"
        else:
            try:
                raw = self._fetch_bytes(f"sites/{site_id}/drives/{drive_id}/items/{item['id']}/content")
                content = _extract_text(extension, raw)
                metadata["content_available"] = True
            except Exception:
                # One unreadable file must not take down the rest of the
                # library's listing -- it is shown, honestly, as unavailable.
                logger.exception("SharePoint content extraction failed for item %s", item.get("id"))
                metadata["content_available"] = False
        return SourceDocument(
            source_id=item["id"], tenant_id=context.tenant_id, project_id=project_id,
            title=name, content=content, web_url=item.get("webUrl"),
            modified_at=datetime.fromisoformat(item["lastModifiedDateTime"])
            if item.get("lastModifiedDateTime") else None,
            metadata=metadata)

    def calendar_events(self, *, start: str, end: str) -> list[dict[str, Any]]:
        params = {"startDateTime": start, "endDateTime": end,
                  "$select": "id,subject,start,end,location,organizer,webLink",
                  "$orderby": "start/dateTime"}
        return [item for page in self._pages("me/calendarView", params) for item in page.values]

    def recent_messages(self, *, limit: int = 25) -> list[dict[str, Any]]:
        params = {"$top": str(min(limit, 100)),
                  "$select": "id,subject,from,receivedDateTime,importance,webLink,bodyPreview",
                  "$orderby": "receivedDateTime desc"}
        return [item for page in self._pages("me/messages", params) for item in page.values][:limit]

    def sharepoint_documents(self, *, site_id: str, drive_id: str,
                             context: UserContext, project_id: str,
                             root_path: str | None = None,
                             folder_item_id: str = "root") -> list[SourceDocument]:
        if project_id not in context.project_ids:
            raise PermissionError(f"User is not authorized for project {project_id!r}")
        if root_path:
            if folder_item_id != "root":
                logger.warning(
                    "SharePoint resource for project %s configures both root_path and "
                    "folder_item_id; root_path takes precedence", project_id)
            item_path = self._resolve_root_path(site_id=site_id, drive_id=drive_id,
                                                root_path=root_path)
        else:
            # folder_item_id is deprecated in favor of root_path (D1): a
            # derived folder name is a guess that fails silently into the
            # wrong client's folder, so this fallback is accepted for one
            # release and logged every time it is actually used.
            logger.warning(
                "SharePoint resource for project %s has no root_path configured; "
                "falling back to the deprecated folder_item_id %r", project_id, folder_item_id)
            item_path = "root" if folder_item_id == "root" else f"items/{folder_item_id}"
        return [self._build_document(item=item, site_id=site_id, drive_id=drive_id,
                                     context=context, project_id=project_id, folder_path=folder_path)
                for item, folder_path in self._walk(
                    site_id=site_id, drive_id=drive_id, item_path=item_path,
                    folder_path="", depth=1)]
