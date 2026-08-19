"""A document must be complete, internally consistent, and executable.

These cover the four things that made assembled agreements not-quite-documents:
cross-references pointing at the wrong section, no parties or signature blocks,
no Schedule A, and clause language that is unenforceable in the governing
jurisdiction.
"""

from tessera_os.clauses import ClauseLibrary, DealProfile, Party

LIBRARY = "fixtures/clause_library"


def library() -> ClauseLibrary:
    return ClauseLibrary.load(LIBRARY)


def parties() -> list[Party]:
    return [
        Party(name="Tessera Holdings LLC", role="member",
              entity_form="a Delaware limited liability company",
              signatory_name="Derrick Carlisle", signatory_title="Managing Partner",
              capital_contribution=1_500_000, units=1500),
        Party(name="Harbor Partners LLC", role="member",
              entity_form="an Oklahoma limited liability company",
              signatory_name="J. Rivera", signatory_title="Manager",
              capital_contribution=500_000, units=500),
    ]


def profile(**overrides) -> DealProfile:
    base = {"opportunity": "Harbor HoldCo", "agreement_type": "operating_agreement",
            "industry": "real_estate", "jurisdiction": "the State of Delaware",
            "counterparty": "Harbor Partners LLC", "fee_at_risk": 0,
            "tessera_capital_at_risk": True, "effective_date": "1 September 2026",
            "parties": parties()}
    base.update(overrides)
    return DealProfile(**base)


# --- cross-references -------------------------------------------------------

def test_cross_references_resolve_to_the_right_section():
    """The indemnity once pointed at the independent-contractor clause."""
    draft = library().assemble(DealProfile(
        opportunity="Green Valley", agreement_type="finders_fee", industry="regulated",
        jurisdiction="the State of Oklahoma", counterparty="Green Valley Holdings LLC",
        counterparty_represented=False, fee_at_risk=250_000))
    numbers = draft.section_numbers()
    body = draft.to_markdown()
    expected = numbers["limitation-liability"]
    assert f"limitation of liability in Section {expected}" in body
    assert numbers["indemnification"] != expected


def test_no_document_ships_with_a_dangling_reference():
    lib = library()
    for agreement_type in ("consulting", "advisory", "finders_fee", "deal_memo",
                           "operating_agreement", "jv", "nda"):
        draft = lib.assemble(profile(agreement_type=agreement_type, parties=[]))
        assert not draft.broken_references(), agreement_type


def test_a_dangling_reference_is_reported_not_hidden():
    lib = library()
    draft = lib.assemble(profile())
    item = draft.selections[0]
    item.variant.text += " See Section {ref:not-a-real-clause}."
    assert draft.broken_references()
    assert "Broken cross-references" in draft.to_markdown()
    assert "[MISSING CLAUSE]" in draft.to_markdown()


# --- parties, schedule, signatures ------------------------------------------

def test_the_document_has_a_preamble_and_signature_blocks():
    body = library().assemble(profile()).to_markdown()
    assert "is entered into as of 1 September 2026" in body
    assert "Tessera Holdings LLC" in body
    assert "## Signatures" in body
    assert "Name: Derrick Carlisle" in body
    assert "Title: Managing Partner" in body


def test_schedule_a_totals_capital_and_units():
    body = library().assemble(profile()).to_markdown()
    assert "## Schedule A — Members, Capital Contributions, and Units" in body
    assert "$1,500,000" in body
    assert "**$2,000,000**" in body
    assert "**2,000**" in body


def test_a_document_without_parties_still_assembles():
    """Clause review does not require signatories."""
    body = library().assemble(profile(parties=[])).to_markdown()
    assert "## Signatures" not in body
    assert "## Schedule A" not in body


def test_a_referenced_attachment_that_is_not_generated_is_flagged():
    """Clauses point at "Schedule A"; without parties there is no Schedule A."""
    without = library().assemble(profile(parties=[]))
    assert "Schedule A" in without.missing_attachments()
    assert "Attachments still to be prepared" in without.to_markdown()

    with_parties = library().assemble(profile())
    assert "Schedule A" not in with_parties.missing_attachments()


def test_exhibits_referenced_by_engagement_paper_are_flagged():
    draft = library().assemble(DealProfile(
        opportunity="RiverBend", agreement_type="consulting", industry="real_estate",
        jurisdiction="the State of Texas", counterparty="RiverBend Residential",
        fee_at_risk=42_000))
    missing = draft.missing_attachments()
    assert "Exhibit A" in missing
    assert "Exhibit B" in missing


# --- jurisdiction overlays --------------------------------------------------

def test_a_variant_unavailable_in_a_jurisdiction_is_not_used():
    """California voids most employee non-solicits."""
    lib = library()
    ids = {i.clause.id for i in lib.assemble(DealProfile(
        opportunity="X", agreement_type="consulting", industry="real_estate",
        jurisdiction="the State of California", counterparty="Y",
        fee_at_risk=42_000)).selections}
    assert "non-solicitation" not in ids


def test_jurisdiction_falls_back_to_an_available_variant():
    lib = library()
    draft = lib.assemble(DealProfile(
        opportunity="X", agreement_type="consulting", industry="real_estate",
        jurisdiction="the State of Oklahoma", counterparty="Y", fee_at_risk=42_000))
    non_solicit = next(i for i in draft.selections if i.clause.id == "non-solicitation")
    assert non_solicit.variant.id == "nonsolicit-mutual-12"


def test_a_jurisdiction_note_reaches_the_counsel_checklist():
    lib = library()
    notes = lib.assemble(DealProfile(
        opportunity="X", agreement_type="consulting", industry="real_estate",
        jurisdiction="the State of Texas", counterparty="Y",
        fee_at_risk=42_000)).counsel_notes()
    assert any("Texas" in note and "reasonable in duration" in note for note in notes)


# --- filled output ----------------------------------------------------------

FILL = {"survival_sections": "3, 5, 7 and 9",
        "company_purpose": "acquiring and operating the Property",
        "manager_name": "Tessera Holdings LLC", "promote_holder": "Tessera Holdings LLC",
        "entity_statute": "the Delaware Limited Liability Company Act",
        "jurisdiction": "the State of Delaware"}


def test_a_filled_document_drops_the_open_terms_list():
    """Substituting into that list turns it into a roll-call of values."""
    lib = library()
    filled = lib.fill(lib.assemble(profile()), dict(FILL))
    assert "Terms still to be filled in" not in filled.markdown
    assert "## Terms supplied" in filled.markdown
    assert "| Manager | Tessera Holdings LLC |" in filled.markdown


def test_a_filled_document_reads_cleanly():
    lib = library()
    filled = lib.fill(lib.assemble(profile()), dict(FILL))
    assert "{" not in filled.markdown
    assert "the the " not in filled.markdown


# --- Word output ------------------------------------------------------------

def test_the_document_renders_to_branded_word(tmp_path):
    from tessera_os.agreement_docx import render_agreement_docx

    lib = library()
    draft = lib.assemble(profile())
    filled = lib.fill(draft, dict(FILL))
    output = render_agreement_docx(draft, tmp_path / "agreement.docx",
                                   markdown=filled.markdown)
    assert output.is_file()
    assert output.stat().st_size > 10_000

    from docx import Document

    rendered = Document(output)
    body = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    header = rendered.sections[0].header.paragraphs[0].text
    footer = rendered.sections[0].footer.paragraphs[0].text
    assert "Operating Agreement — Harbor HoldCo" in body
    assert "DRAFT — FOR QUALIFIED COUNSEL REVIEW" in body
    assert "TESSERA GROUP" in header
    assert "Harbor HoldCo" in footer
    assert rendered.tables


def test_a_structure_memo_renders_on_the_same_paper(tmp_path):
    """Same brand, different banner -- a memo is not a draft agreement."""
    from docx import Document

    from tessera_os.agreement_docx import render_structure_docx
    from tessera_os.governance import VentureProfile, recommend_structure

    rec = recommend_structure(VentureProfile(
        venture="Harbor HoldCo", home_state="Oklahoma", real_property=True))
    output = render_structure_docx(rec, tmp_path / "memo.docx")
    rendered = Document(output)
    header = rendered.sections[0].header.paragraphs[0].text
    body = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "TESSERA GROUP" in header
    assert "not legal or tax advice" in header
    assert "Structure Recommendation — Harbor HoldCo" in body
    # Italic and third-level headings must not reach the page as raw markdown.
    assert "*" not in body
    assert "###" not in body
