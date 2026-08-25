"""Phase 2 (revised): library-shaped document reading.

Covers every acceptance criterion in
docs/BUILD_BRIEF_PHASE_2_LIBRARY_READING.md section 4 -- path-based project
resources, the depth-bounded recursive walk, lifecycle derived from folder
(never Basis, per D2), dropping the ProjectId column, content extraction from
the file, and explicit ACLs derived from the trust zone.
"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument

from tessera_os.identity import PARTNER_GROUP, ZoneAccessError, ZonePolicy
from tessera_os.integrations import IntegrationError, MicrosoftGraphReader
from tessera_os.knowledge import KnowledgeIndex
from tessera_os.microsoft import (
    AllowlistedSharePointReader,
    MicrosoftConfigurationError,
    MicrosoftPilotSettings,
    SharePointProjectResource,
)
from tessera_os.schemas import SourceDocument, UserContext

BASE = MicrosoftGraphReader.base_url


def resource(**overrides) -> SharePointProjectResource:
    defaults = {"site_id": "site-x", "drive_id": "drive-x",
                "root_path": "Projects/Internal Pilot", "zone": "internal"}
    defaults.update(overrides)
    return SharePointProjectResource(**defaults)


def settings_for(project_id: str = "internal-pilot", **overrides) -> MicrosoftPilotSettings:
    return MicrosoftPilotSettings(
        enabled=True, tenant_id="t", client_id="c", client_secret="s", cache_key="k",
        project_resources={project_id: resource(**overrides)})


def context(project_ids=("internal-pilot",), groups=("tessera_partner",)) -> UserContext:
    return UserContext(tenant_id="t", user_id="derrick",
                       project_ids=set(project_ids), group_ids=set(groups))


def reader_for(settings, transport, content_transport=None) -> AllowlistedSharePointReader:
    if content_transport is None:
        def content_transport(_url, _headers):
            return b""
    return AllowlistedSharePointReader(
        settings=settings, token_provider=lambda _user_id: "token",
        graph_factory=lambda provider: MicrosoftGraphReader(
            provider, transport=transport, content_transport=content_transport))


def docx_bytes(*paragraphs: str) -> bytes:
    document = DocxDocument()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- 2.1 -- path-based project resources -------------------------------------------------------

def test_root_path_resolves_via_path_addressing():
    calls = []

    def transport(url, headers):
        calls.append(url)
        if "root:/" in url:
            assert url == f"{BASE}/sites/site-x/drives/drive-x/root:/Projects/Internal%20Pilot:"
            return {"id": "folder-1", "folder": {}}
        assert "items/folder-1/children" in url
        return {"value": []}

    reader = reader_for(settings_for(), transport)
    assert reader.project_documents(context=context(), project_id="internal-pilot") == []
    assert any("root:/" in call for call in calls)


def test_a_root_path_that_does_not_resolve_raises_a_configuration_error_naming_the_path():
    def transport(url, headers):
        raise IntegrationError("not found")

    reader = reader_for(settings_for(root_path="Projects/Missing"), transport)
    with pytest.raises(MicrosoftConfigurationError, match="Projects/Missing"):
        reader.project_documents(context=context(), project_id="internal-pilot")


def test_root_path_rejects_traversal_and_a_leading_slash():
    from pydantic import ValidationError

    with pytest.raises(ValidationError, match="root_path"):
        SharePointProjectResource(site_id="s", drive_id="d", root_path="/absolute/path")
    with pytest.raises(ValidationError, match="root_path"):
        SharePointProjectResource(site_id="s", drive_id="d", root_path="Projects/../Secret")


def test_both_root_path_and_folder_item_id_present_root_path_wins_with_a_warning(caplog):
    calls = []

    def transport(url, headers):
        calls.append(url)
        if "root:/" in url:
            return {"id": "correct-folder", "folder": {}}
        return {"value": []}

    settings = settings_for(root_path="Projects/Internal Pilot", folder_item_id="legacy-folder")
    reader = reader_for(settings, transport)
    with caplog.at_level("WARNING"):
        reader.project_documents(context=context(), project_id="internal-pilot")
    assert any("root_path takes precedence" in message for message in caplog.messages)
    assert any("items/correct-folder/children" in call for call in calls)
    assert not any("legacy-folder" in call for call in calls)


def test_deprecated_folder_item_id_alone_still_works_and_logs(caplog):
    def transport(url, headers):
        assert "items/legacy-folder/children" in url
        return {"value": []}

    settings = settings_for(root_path=None, folder_item_id="legacy-folder")
    reader = reader_for(settings, transport)
    with caplog.at_level("WARNING"):
        documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert documents == []
    assert any("deprecated folder_item_id" in message for message in caplog.messages)


# --- 2.2 -- the recursive walk ------------------------------------------------------------------

def test_recursive_walk_returns_nested_files_and_never_returns_folders(caplog):
    tree = {
        "root-item": [{"id": "l1", "name": "L1", "folder": {}}],
        "l1": [{"id": "l2", "name": "L2", "folder": {}}],
        "l2": [{"id": "l3", "name": "L3", "folder": {}}],
        "l3": [{"id": "l4", "name": "L4", "folder": {}}],
        "l4": [
            {"id": "deep-5", "name": "depth5.txt", "file": {}, "size": 5},
            {"id": "l5", "name": "L5", "folder": {}},
        ],
        "l5": [{"id": "deep-6", "name": "depth6.txt", "file": {}, "size": 5}],
    }
    requested = []

    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        for key, children in tree.items():
            if f"items/{key}/children" in url:
                requested.append(key)
                return {"value": children}
        raise AssertionError(f"unexpected url {url}")

    reader = reader_for(settings_for(), transport)
    with caplog.at_level("WARNING"):
        documents = reader.project_documents(context=context(), project_id="internal-pilot")

    titles = {document.title for document in documents}
    assert titles == {"depth5.txt"}  # depth 5 is returned, depth 6 never is; folders never appear
    assert "l5" not in requested  # a folder beyond depth 5 is never even walked
    assert any("skipping" in message for message in caplog.messages)


def test_paging_is_followed_at_every_level_not_just_the_first():
    counts = {"root-item": 0, "sub": 0}

    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        if "items/root-item/children" in url:
            counts["root-item"] += 1
            if counts["root-item"] == 1:
                return {"value": [{"id": "sub", "name": "Sub", "folder": {}}],
                        "@odata.nextLink":
                            f"{BASE}/sites/site-x/drives/drive-x/items/root-item/children?page=2"}
            return {"value": [{"id": "a", "name": "a.txt", "file": {}, "size": 1}]}
        if "items/sub/children" in url:
            counts["sub"] += 1
            if counts["sub"] == 1:
                return {"value": [{"id": "b", "name": "b.txt", "file": {}, "size": 1}],
                        "@odata.nextLink":
                            f"{BASE}/sites/site-x/drives/drive-x/items/sub/children?page=2"}
            return {"value": [{"id": "c", "name": "c.txt", "file": {}, "size": 1}]}
        raise AssertionError(f"unexpected url {url}")

    reader = reader_for(settings_for(), transport, content_transport=lambda url, headers: b"x")
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert {document.title for document in documents} == {"a.txt", "b.txt", "c.txt"}


# --- 2.3 -- lifecycle from folder, per D2 -------------------------------------------------------

def test_lifecycle_is_derived_from_the_first_folder_below_root_path():
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        if "items/root-item/children" in url:
            return {"value": [
                {"id": "approved-folder", "name": "Approved", "folder": {}},
                {"id": "drafts-folder", "name": "Drafts", "folder": {}},
                {"id": "source-folder", "name": "Source", "folder": {}},
                {"id": "misc-folder", "name": "Misc", "folder": {}},
                {"id": "top", "name": "top.txt", "file": {}, "size": 3},
            ]}
        if "items/approved-folder/children" in url:
            return {"value": [{"id": "memo", "name": "memo.docx", "file": {}, "size": 3}]}
        if "items/drafts-folder/children" in url:
            return {"value": [{"id": "draft1", "name": "draft.docx", "file": {}, "size": 3}]}
        if "items/source-folder/children" in url:
            return {"value": [
                {"id": "fin", "name": "client-financials.xlsx", "file": {}, "size": 3}]}
        if "items/misc-folder/children" in url:
            return {"value": [{"id": "note", "name": "note.docx", "file": {}, "size": 3}]}
        raise AssertionError(f"unexpected url {url}")

    reader = reader_for(settings_for(), transport, content_transport=lambda url, headers: b"text")
    documents = {document.title: document
                for document in reader.project_documents(context=context(),
                                                          project_id="internal-pilot")}

    assert documents["memo.docx"].metadata["lifecycle"] == "approved"
    assert documents["draft.docx"].metadata["lifecycle"] == "draft"
    assert documents["client-financials.xlsx"].metadata["lifecycle"] == "source"
    assert documents["note.docx"].metadata["lifecycle"] == "source"
    assert documents["top.txt"].metadata["lifecycle"] == "source"  # no lifecycle folder
    assert documents["memo.docx"].metadata["folder_path"] == "Approved"


def test_a_source_document_never_carries_a_basis_because_it_is_evidence_not_a_position():
    """D2: a client's uploaded file is evidence, not a Tessera position. Writing
    Source/ -> synthetic_reference would stamp a real client document as
    invented, inside the exact field the review queue and the adoption ledger
    read to decide what is real -- a provenance violation, not a convenience.
    """
    document = SourceDocument(source_id="x", tenant_id="t", project_id="p", title="a.txt",
                              content="", metadata={"lifecycle": "source"})
    assert not hasattr(document, "basis")
    assert "basis" not in type(document).model_fields


# --- 2.4 -- drop the ProjectId column requirement -----------------------------------------------

def test_an_untagged_file_inside_the_project_folder_is_returned():
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc", "name": "untagged.txt", "file": {}, "size": 3}]}

    reader = reader_for(settings_for(), transport, content_transport=lambda url, headers: b"x")
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert documents[0].title == "untagged.txt"
    assert documents[0].project_id == "internal-pilot"


def test_a_stale_projectid_column_cannot_override_the_path():
    """The column is never read at all now (D1) -- a file carries whatever a
    leftover SharePoint list column claims, and it changes nothing."""
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{
            "id": "doc", "name": "spoofed.txt", "file": {}, "size": 3,
            "listItem": {"fields": {"ProjectId": "some-other-project"}},
        }]}

    reader = reader_for(settings_for(), transport, content_transport=lambda url, headers: b"x")
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert documents[0].project_id == "internal-pilot"


def test_a_file_under_a_different_projects_folder_is_never_even_requested():
    """Scope is the path (D1): a project's reader resolves only its own
    root_path, so another project's subtree is never walked -- there is
    nothing for a stale column to override."""
    requested = []

    def transport(url, headers):
        if "root:/Projects/Alpha:" in url:
            return {"id": "alpha-root", "folder": {}}
        if "root:/Projects/Bravo:" in url:
            return {"id": "bravo-root", "folder": {}}
        requested.append(url)
        if "items/alpha-root/children" in url:
            return {"value": [{"id": "a", "name": "alpha.txt", "file": {}, "size": 1}]}
        if "items/bravo-root/children" in url:
            return {"value": [{"id": "b", "name": "bravo.txt", "file": {}, "size": 1}]}
        raise AssertionError(f"unexpected url {url}")

    settings = MicrosoftPilotSettings(
        enabled=True, tenant_id="t", client_id="c", client_secret="s", cache_key="k",
        project_resources={
            "alpha": resource(root_path="Projects/Alpha", zone="internal"),
            "bravo": resource(root_path="Projects/Bravo", zone="internal"),
        })
    reader = reader_for(settings, transport, content_transport=lambda url, headers: b"x")

    documents = reader.project_documents(
        context=context(project_ids=("alpha", "bravo")), project_id="alpha")
    assert {document.title for document in documents} == {"alpha.txt"}
    assert not any("bravo-root" in call for call in requested)


def test_the_golden_rule_boundary_still_reads_project_id_off_the_document():
    """ZonePolicy.check_citation is the golden-rule boundary this phase must
    not regress: a document from one engagement's zone never cites into
    another's artifact, regardless of what a stale column might claim."""
    policy = ZonePolicy(resource_zones={"alpha": "engagement", "bravo": "engagement"},
                        resource_clients={"alpha": "client-a", "bravo": "client-b"})
    with pytest.raises(ZoneAccessError, match="outside"):
        policy.check_citation(source_project_id="alpha", artifact_project_id="bravo",
                              artifact_client_id="client-b")


# --- 2.5 -- content extraction, per D3 ----------------------------------------------------------

def test_docx_content_is_extracted_as_paragraph_text():
    raw = docx_bytes("First paragraph.", "Second paragraph.")

    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc", "name": "memo.docx", "file": {}, "size": len(raw)}]}

    reader = reader_for(settings_for(), transport, content_transport=lambda url, headers: raw)
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert "First paragraph." in documents[0].content
    assert "Second paragraph." in documents[0].content
    assert documents[0].metadata["content_available"] is True


def test_unsupported_extension_lists_with_no_content():
    fetched = []

    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc", "name": "model.xlsx", "file": {}, "size": 10}]}

    def content_transport(url, headers):
        fetched.append(url)
        return b""

    reader = reader_for(settings_for(), transport, content_transport=content_transport)
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert documents[0].content == ""
    assert documents[0].metadata["content_available"] is False
    assert fetched == []  # an unsupported type is never even fetched


def test_a_file_over_10mb_is_listed_but_not_fetched():
    fetched = []

    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [
            {"id": "doc", "name": "big.txt", "file": {}, "size": 10 * 1024 * 1024 + 1}]}

    def content_transport(url, headers):
        fetched.append(url)
        return b""

    reader = reader_for(settings_for(), transport, content_transport=content_transport)
    documents = reader.project_documents(context=context(), project_id="internal-pilot")
    assert documents[0].content == ""
    assert documents[0].metadata["content_available"] is False
    assert documents[0].metadata["skipped_reason"] == "size"
    assert fetched == []


def test_extraction_failure_on_one_file_does_not_abort_the_walk(caplog):
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [
            {"id": "broken", "name": "broken.txt", "file": {}, "size": 3},
            {"id": "ok", "name": "ok.txt", "file": {}, "size": 3},
        ]}

    def content_transport(url, headers):
        if "items/broken/content" in url:
            raise IntegrationError("boom")
        return b"fine"

    reader = reader_for(settings_for(), transport, content_transport=content_transport)
    with caplog.at_level("ERROR"):
        documents = {document.title: document
                    for document in reader.project_documents(context=context(),
                                                              project_id="internal-pilot")}
    assert documents["broken.txt"].metadata["content_available"] is False
    assert documents["ok.txt"].content == "fine"
    assert documents["ok.txt"].metadata["content_available"] is True


# --- 2.6 -- explicit ACLs, per D4 ----------------------------------------------------------------

def test_an_internal_zone_document_carries_the_partner_group():
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc", "name": "memo.txt", "file": {}, "size": 3}]}

    reader = reader_for(settings_for(zone="internal"), transport,
                        content_transport=lambda url, headers: b"x")
    documents = reader.project_documents(
        context=context(groups=("tessera_partner",)), project_id="internal-pilot")
    assert documents[0].allowed_group_ids == {PARTNER_GROUP}


def test_an_engagement_zone_document_carries_its_own_group_not_the_partner_group_alone():
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc", "name": "memo.txt", "file": {}, "size": 3}]}

    settings = settings_for(project_id="acme-workspace", zone="engagement",
                            client_id="client-acme", root_path="Projects/Acme")
    reader = reader_for(settings, transport, content_transport=lambda url, headers: b"x")
    documents = reader.project_documents(
        context=UserContext(tenant_id="t", user_id="u", project_ids={"acme-workspace"},
                            group_ids={"engagement:client-acme"}),
        project_id="acme-workspace")
    assert documents[0].allowed_group_ids == {"engagement:client-acme"}
    assert PARTNER_GROUP not in documents[0].allowed_group_ids


def test_a_document_from_the_portal_listing_is_also_findable_via_knowledge_search():
    """The divergence that made every SharePoint document invisible to
    knowledge search while still listed in the portal (see the evidence in
    docs/BUILD_BRIEF_PHASE_2_LIBRARY_READING.md section 2.3) is a defect a
    per-path test would not catch -- exercise both paths from one reader
    output."""
    def transport(url, headers):
        if "root:/" in url:
            return {"id": "root-item", "folder": {}}
        return {"value": [{"id": "doc-1", "name": "memo.txt", "file": {}, "size": 20}]}

    reader = reader_for(settings_for(zone="internal"), transport,
                        content_transport=lambda url, headers: b"Milestone is current")
    partner = context(groups=("tessera_partner",))
    documents = reader.project_documents(context=partner, project_id="internal-pilot")
    assert len(documents) == 1  # visible on the portal-listing path

    index = KnowledgeIndex()
    index.ingest(documents)
    hits = index.search("milestone", context=partner, project_id="internal-pilot")
    assert [hit.source_id for hit in hits] == ["doc-1"]  # and on the knowledge-search path
